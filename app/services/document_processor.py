import os
import tempfile
from typing import List, Dict, Any
from pathlib import Path
from datetime import datetime

import PyPDF2
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument

from app.config import settings

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.MAX_CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    async def process_uploaded_file(self, file_content: bytes, filename: str) -> List[LangchainDocument]:
        """Process uploaded file and return list of documents"""
        # Save file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as tmp_file:
            tmp_file.write(file_content)
            tmp_file_path = tmp_file.name
        
        try:
            # Extract text based on file type
            if filename.lower().endswith('.pdf'):
                text = self._extract_pdf_text(tmp_file_path)
            elif filename.lower().endswith('.docx'):
                text = self._extract_docx_text(tmp_file_path)
            else:
                raise ValueError(f"Unsupported file type: {filename}")
            
            # Create document with metadata
            document = LangchainDocument(
                page_content=text,
                metadata={
                    "source": filename,
                    "file_type": Path(filename).suffix,
                    "processed_at": str(datetime.now())  # CHANGED THIS LINE
                }
            )
            
            # Split into chunks
            chunks = self.text_splitter.split_documents([document])
            
            return chunks
            
        finally:
            # Clean up temporary file
            os.unlink(tmp_file_path)
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    
    def process_text(self, text: str, metadata: Dict[str, Any] = None) -> List[LangchainDocument]:
        """Process raw text and return chunks"""
        if metadata is None:
            metadata = {}
            
        document = LangchainDocument(
            page_content=text,
            metadata=metadata
        )
        
        return self.text_splitter.split_documents([document])